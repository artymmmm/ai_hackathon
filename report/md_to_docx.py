"""Конвертация отчётов из markdown в .docx — задание требует .pdf или .docx.

Намеренно простой конвертер: заголовки, абзацы, списки, таблицы, блоки кода. Ничего сверх
того, что реально встречается в наших отчётах.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Pt, RGBColor

BOLD = re.compile(r"\*\*(.+?)\*\*")
CODE = re.compile(r"`([^`]+)`")


LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_link(par, text: str, url: str, bold: bool = False) -> None:
    """Гиперссылка: python-docx не умеет их из коробки, собираем элемент руками."""
    rel_id = par.part.relate_to(url, RT.HYPERLINK, is_external=True)
    run = par.add_run(text)
    run.bold = bold
    run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    run.font.underline = True
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rel_id)
    par._p.append(link)
    link.append(run._element)


def add_rich(par, text: str) -> None:
    """Разбирает **жирный**, `код` и [ссылку](адрес) внутри строки."""
    pos = 0
    for m in re.finditer(r"\*\*(.+?)\*\*|`([^`]+)`|\[([^\]]+)\]\(([^)]+)\)", text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        if m.group(1) is not None:
            inner = LINK.fullmatch(m.group(1))
            if inner:
                add_link(par, inner.group(1), inner.group(2), bold=True)
            else:
                par.add_run(m.group(1)).bold = True
        elif m.group(2) is not None:
            r = par.add_run(m.group(2))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x8B, 0x1A, 0x1A)
        else:
            add_link(par, m.group(3), m.group(4))
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])


def convert(md_path: str, docx_path: str) -> None:
    lines = Path(md_path).read_text(encoding="utf-8").splitlines()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            i += 1
            buf = []
            while i < len(lines) and not lines[i].startswith("```"):
                buf.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(buf))
            r.font.name = "Consolas"
            r.font.size = Pt(8.5)
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].replace("|", "").strip()) <= set("-: "):
            header = [c.strip() for c in line.strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            t = doc.add_table(rows=1, cols=len(header))
            t.style = "Light Grid Accent 1"
            for c, h in zip(t.rows[0].cells, header):
                c.text = ""
                add_rich(c.paragraphs[0], h)
                for run in c.paragraphs[0].runs:
                    run.bold = True
            for row in rows:
                cells = t.add_row().cells
                for c, val in zip(cells, row):
                    c.text = ""
                    add_rich(c.paragraphs[0], val)
            doc.add_paragraph()
            continue

        if m := re.match(r"^(#{1,4})\s+(.*)", line):
            doc.add_heading(m.group(2).replace("**", ""), level=len(m.group(1)))
        elif m := re.match(r"^\s*(?:([-*])|\d+\.)\s+(.*)", line):
            # продолжение пункта свёрстано с отступом на следующих строках — склеиваем,
            # иначе хвост пункта уезжает в отдельный абзац
            buf = [m.group(2)]
            while (i + 1 < len(lines) and lines[i + 1].startswith(" ") and lines[i + 1].strip()
                   and not re.match(r"^\s*(?:[-*]|\d+\.)\s", lines[i + 1])):
                i += 1
                buf.append(lines[i].strip())
            style = "List Bullet" if m.group(1) else "List Number"
            add_rich(doc.add_paragraph(style=style), " ".join(buf))
        elif line.startswith(">"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_rich(p, line.lstrip("> "))
            for r in p.runs:
                r.italic = True
        elif line.strip() == "---":
            doc.add_paragraph("─" * 40)
        elif line.strip():
            # markdown в отчётах свёрстан с жёсткими переносами; склеиваем подряд идущие
            # строки обычного текста в один абзац, иначе docx выглядит рваным
            buf = [line.strip()]
            while (i + 1 < len(lines) and lines[i + 1].strip()
                   and not re.match(r"^(#{1,4}\s|\s*[-*]\s|\s*\d+\.\s|\||```|>|---$)", lines[i + 1])):
                i += 1
                buf.append(lines[i].strip())
            add_rich(doc.add_paragraph(), " ".join(buf))
        i += 1

    doc.save(docx_path)
    print(f"{md_path} -> {docx_path}")


if __name__ == "__main__":
    for md in sys.argv[1:]:
        convert(md, str(Path(md).with_suffix(".docx")))
